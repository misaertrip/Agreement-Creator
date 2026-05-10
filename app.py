import streamlit as st
from docxtpl import DocxTemplate
from docx import Document
from docx import Document as DocxDoc
from copy import deepcopy
import io
import os
from datetime import date

# Page configuration
st.set_page_config(page_title="Agreement Generator", layout="wide")
st.title("📄 Professional Agreement Generator")

# -----------------------------
# 1. INPUT SECTION
# -----------------------------
col1, col2 = st.columns(2)
with col1:
    org_name = st.text_input("Organisation Name")
    doc_date = st.date_input("Agreement Date", value=date.today())
    doc_type = st.text_input("Document Type")
    doc_number = st.text_input("GSTIN / Document Number")
    email = st.text_input("Client Email")

with col2:
    address = st.text_area("Registered Address")
    org_sign_name = st.text_input("Organisation Signatory Name")
    aer_sign_name = st.text_input("Aertrip Signatory Name")

designation_options = ["Director", "Partner", "Proprietor", "Vice President - Operations", "Other"]
st.write("---")
c3, c4 = st.columns(2)
with c3:
    org_designation = st.selectbox("Organisation Signatory Designation", designation_options)
    if org_designation == "Other":
        org_designation = st.text_input("Enter Custom Org Designation")
with c4:
    aer_designation = st.selectbox("Aertrip Signing Designation", designation_options)
    if aer_designation == "Other":
        aer_designation = st.text_input("Enter Custom Aertrip Designation")

st.write("---")
st.subheader("Annexure Settings")
col5, col6 = st.columns(2)
with col5:
    annexure_a_choice = st.radio("Include Annexure A (Fees)?", ["Yes", "No"], horizontal=True)
with col6:
    annexure_b_choice = st.radio("Include Annexure B (Related Parties)?", ["Yes", "No"], horizontal=True)

party_names = []
if annexure_b_choice == "Yes":
    num_parties = st.number_input("How many related parties?", min_value=1, step=1)
    for i in range(int(num_parties)):
        name = st.text_input(f"Related Party {i+1} Name", key=f"party_{i}")
        if name:
            party_names.append(name)

# -----------------------------
# 2. GENERATION ENGINE
# -----------------------------
if st.button("🚀 Generate & Download Agreement", type="primary"):
    if not org_name or not doc_number:
        st.error("Please fill Organisation Name and Document Number!")
    else:
        try:
            BASE_DIR = os.path.dirname(os.path.abspath(__file__))
            template_path = os.path.join(BASE_DIR, "template.docx")
            ann_a_path = os.path.join(BASE_DIR, "annexure_a.docx")

            doc = DocxTemplate(template_path)

            ann_a_line = "The implications of Aertrip Fees are outlined in Annexure A." if annexure_a_choice == "Yes" else ""
            ann_b_line = "and its related entities as mentioned in Annexure B" if annexure_b_choice == "Yes" else ""

            context = {
                "org_name": org_name,
                "date": doc_date.strftime('%d %B %Y'),
                "document_type": doc_type,
                "document_number": doc_number,
                "address": address,
                "email": email,
                "org_sign_name": org_sign_name,
                "org_sign_designation": org_designation,
                "aer_sign_name": aer_sign_name,
                "aer_sign_designation": aer_designation,
                "annexure_a_line": ann_a_line,
                "annexure_b_line": ann_b_line
            }

            # --- Annexure A ---
            if annexure_a_choice == "Yes" and os.path.exists(ann_a_path):
                sub_a_doc = Document()
                sub_a_doc.add_page_break()
                source_a = Document(ann_a_path)
                for element in source_a.element.body:
                    if element.tag.endswith('}sectPr'):
                        continue
                    sub_a_doc.element.body.append(deepcopy(element))

                a_buf = io.BytesIO()
                sub_a_doc.save(a_buf)
                a_buf.seek(0)
                context["annexure_a_section"] = doc.new_subdoc(a_buf)
            else:
                context["annexure_a_section"] = ""

            # --- Annexure B ---
            if annexure_b_choice == "Yes" and party_names:
                sub_b_doc = Document()
                sub_b_doc.add_page_break()
                sub_b_doc.add_heading("ANNEXURE B - CLIENT'S ENTITIES", level=1)
                for i, name in enumerate(party_names, 1):
                    sub_b_doc.add_paragraph(f"{i}. {name}")

                b_buf = io.BytesIO()
                sub_b_doc.save(b_buf)
                b_buf.seek(0)
                context["annexure_b_section"] = doc.new_subdoc(b_buf)
            else:
                context["annexure_b_section"] = ""

            # --- Render ---
            doc.render(context)

            # --- Page Numbering Fix ---
            temp_buf = io.BytesIO()
            doc.save(temp_buf)
            temp_buf.seek(0)

            clean = DocxDoc(temp_buf)
            ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

            for sectPr in clean.element.body.iter(f"{{{ns}}}sectPr"):
                for pgNumType in sectPr.findall(f"{{{ns}}}pgNumType"):
                    start_attr = f"{{{ns}}}start"
                    if start_attr in pgNumType.attrib:
                        del pgNumType.attrib[start_attr]

            final_buffer = io.BytesIO()
            clean.save(final_buffer)
            final_buffer.seek(0)

            st.success("✅ Agreement generated successfully!")
            st.download_button(
                label="📥 Download Word File",
                data=final_buffer,
                file_name=f"Agreement_{org_name}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        except Exception as e:
            st.error(f"Error: {str(e)}")
